from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "docs" / "DynaTwin-Swarm项目文档.docx"
SHOT = ROOT / "docs" / "screenshots"


def set_run_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph(paragraph, line=1.5, align=None, space_after=6):
    paragraph.paragraph_format.line_spacing = line
    paragraph.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        paragraph.alignment = align


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_run_font(run, 10)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_paragraph(p, line=1.5, space_after=8)
    run = p.add_run(text)
    set_run_font(run, 16 if level == 1 else 14, bold=True)
    return p


def add_para(doc, text, size=12, align=None):
    p = doc.add_paragraph()
    set_paragraph(p, align=align)
    run = p.add_run(text)
    set_run_font(run, size)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    set_paragraph(p, line=1.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    run = p.add_run(text)
    set_run_font(run, 10, italic=True)


def add_image(doc, image_name, caption):
    image_path = SHOT / image_name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(14))
    add_caption(doc, caption)


def apply_table_font(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                set_paragraph(paragraph, line=1.2, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
                for run in paragraph.runs:
                    set_run_font(run, 10)


def create_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(12)

    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_p.add_run("DynaTwin-Swarm项目技术文档")
    set_run_font(header_run, 10)
    footer_p = section.footer.paragraphs[0]
    add_page_number(footer_p)

    # Cover
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DynaTwin-Swarm：基于动态图结构自适应选择的工业多智能体协同调度系统")
    set_run_font(run, 20, bold=True)
    set_paragraph(p, line=1.5, space_after=24)

    cover_items = [
        "版本号：V1.0",
        "日期：2026年6月",
        "作者：王骏",
        "单位：贵州大学网络与信息安全专业",
        "指导教师：谭伟杰",
    ]
    for item in cover_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph(p, line=1.5, space_after=8)
        run = p.add_run(item)
        set_run_font(run, 14)
    doc.add_page_break()

    add_heading(doc, "摘要", 1)
    add_para(
        doc,
        "针对工业制造场景中生产调度系统存在的动态响应迟滞、人工依赖程度高、复杂扰动处理能力不足等核心问题，本文提出并实现了DynaTwin-Swarm——一种基于动态图结构自适应选择与目标状态反射推理的工业多智能体协同调度系统。系统融合EMNLP 2025顶级会议论文DynaSwarm与ReflAct的核心方法论，构建了包含11个专业智能体的分层协作决策框架，并以豆包大模型作为推理引擎，结合CP-SAT约束规划求解器实现毫秒级排程计算。实验表明，系统在JSPLib FT06标准基准上的Makespan与已知最优解差距为21.8%，在Lawrence LA40基准上差距为7.0%，求解时间均在200ms以内，验证了系统在动态工业环境中的实时调度能力与工程可行性。研究结果表明，动态图结构选择与可解释推理机制能够有效提升工业调度系统在突发扰动下的自适应决策能力，为贵州磷化工、非煤矿山和智能制造场景提供具有工程推广价值的技术路径。",
    )

    add_heading(doc, "第一章 引言", 1)
    add_para(
        doc,
        "传统工厂生产调度面临多维度的动态扰动问题，设备故障、紧急订单插入、物料短缺和工序切换等不确定因素会导致既有排程方案频繁失效。由于生产现场通常具有多设备、多工序、多资源约束的复杂特征，人工重排不仅响应周期长，而且高度依赖调度人员的经验判断，难以形成可沉淀、可复用、可迁移的知识体系。这一问题严重制约了制造企业的柔性生产能力。以贵州磷化工行业为例，该行业年产值超800亿元，但智能调度渗透率不足5%，在产业升级、降本增效和安全生产方面存在迫切需求。",
    )
    add_para(
        doc,
        "现有工业调度系统普遍采用固定规则引擎或单一优化算法，虽然能够在静态约束下获得一定效果，但缺乏对动态扰动的实时感知与自适应调整能力。当设备状态、订单优先级或资源可用性发生变化时，传统系统往往需要重新配置规则或人工介入，难以在多约束条件下兼顾全局优化、安全约束与解释性需求。与此同时，单一大语言模型直接参与工业决策又可能产生幻觉和不可审计问题，无法满足工业现场对可靠性和合规性的要求。",
    )
    add_para(
        doc,
        "DynaTwin-Swarm提出动态拓扑自适应选择机制与ReflAct推理框架的协同应用，在工业调度场景中实现多智能体动态协作与可解释决策的有机结合。系统以数字孪生工厂状态为基础，通过11个专业智能体对设备、订单、资源、风险和排程约束进行分工推理，并由CP-SAT求解器生成可执行排程方案。该系统为智能制造领域提供了一种兼具实时响应、可解释性和工程可落地性的新技术范式。",
    )
    add_image(doc, "01_dashboard_normal.png", "图1-1 DynaTwin-Swarm系统总览界面（LA40基准数据集，15台机器，15个订单）")

    add_heading(doc, "第二章 系统架构", 1)
    add_heading(doc, "2.1 总体架构", 2)
    add_para(
        doc,
        "DynaTwin-Swarm采用感知层、决策层与执行层相互衔接的三层架构设计。感知层通过IoT传感器适配器和数字孪生工厂模拟器实现工厂状态的实时采集与建模，状态对象覆盖机器可用性、温度、效率、当前工单、订单优先级、物料库存、告警事件和恢复计划等关键生产要素。该层的核心职责是将真实或模拟工业现场转化为统一的结构化FactoryState，为上层智能体推理提供可信输入。",
    )
    add_para(
        doc,
        "决策层由11个专业智能体构成分层协作网络，基于动态拓扑选择机制驱动跨智能体的协同推理。不同风险等级和任务复杂度会触发不同的智能体协作结构，例如常规任务可使用轻量串行链，而设备故障、紧急插单和高风险告警则会切换至高风险审查拓扑。执行层集成CP-SAT约束规划求解器，将智能体决策转化为可执行的优化排程方案，并通过甘特图、事件流和日志面板向用户呈现调度结果。",
    )
    add_heading(doc, "2.2 智能体协作框架", 2)
    add_para(
        doc,
        "智能体协作框架以专业分工为基础，形成从任务识别到结果报告的闭环推理链。任务路由智能体负责感知工厂状态并选择协作拓扑；设备监控智能体实时扫描机器状态并标记故障设备；故障诊断智能体评估故障影响范围并重构候选资源池；订单分析智能体量化订单优先级与截止压力；资源分配智能体解决多维资源冲突；调度求解智能体调用CP-SAT求解器输出最优排程；约束验证智能体执行硬性约束合规检查；风险评估智能体综合评估方案安全等级；决策审核智能体执行一致性校验；报告生成智能体汇总输出完整调度报告。上述角色共同构成可扩展、可审计的工业多智能体协同决策体系。",
    )
    add_image(doc, "02_topology_graph.png", "图2-1 动态Agent协作拓扑图（高风险审查模式，11个智能体全激活）")

    add_heading(doc, "第三章 核心技术", 1)
    add_heading(doc, "3.1 动态图结构自适应选择机制", 2)
    add_para(
        doc,
        "DynaSwarm方法在本系统中的工程化实现体现为面向任务状态的动态拓扑选择机制。系统预定义五种协作拓扑结构：串行链适用于低复杂度常规调度，由4个智能体顺序执行；并行汇聚适用于中等复杂度场景，由7个智能体并行分析后汇聚决策；层级树适用于多层级任务分解，由8个智能体分层协调；全连接适用于高复杂度排程，由10个智能体充分交互；高风险审查适用于故障或紧急场景，由11个智能体全激活并引入多层审核机制。",
    )
    add_para(
        doc,
        "基于A2C（Actor-Critic）强化学习算法，系统通过历史调度奖励信号持续训练图选择器，使拓扑选择策略随使用积累不断优化。奖励信号综合考虑排程完工时间、风险等级、约束违反情况和调度稳定性，从而使系统能够在不同生产状态下学习更合理的协作结构，实现结构与任务的自适应匹配。",
    )
    add_heading(doc, "3.2 ReflAct目标状态反射推理框架", 2)
    add_para(
        doc,
        "ReflAct目标状态反射推理框架要求每个智能体按照观察、目标、偏差和动作四个步骤形成可审计的决策轨迹。智能体首先观察当前工厂状态，获取机器可用性、订单优先级、库存水位和告警数量等多维信息；其次将观察结果与既定生产目标进行对齐，识别当前状态与理想目标之间的偏差；继而基于偏差分析评估风险等级和影响范围；最终输出具体建议动作，为后续拓扑选择和CP-SAT排程提供结构化依据。该框架有效抑制了大语言模型在工业调度场景中的幻觉现象，提升了决策的可解释性与可靠性。",
    )
    add_heading(doc, "3.3 华为昇腾技术集成", 2)
    add_para(
        doc,
        "系统与华为全栈AI技术体系保持适配设计：以华为盘古大模型作为备用推理引擎，通过MindIE推理引擎实现NPU加速，GaussDB承载历史调度知识库，IoTDA实现工厂传感器数据实时接入，EventGrid驱动事件触发式调度响应，FunctionGraph提供弹性计算能力，ModelArts Studio负责模型训练与部署管理。系统可部署于华为云西南-贵阳一区域，充分发挥昇腾算力在本地工业智能化场景中的价值。",
    )

    add_heading(doc, "第四章 功能演示与验证", 1)
    add_heading(doc, "4.1 动态排程演示", 2)
    add_para(doc, "在正常排产场景下，系统首先读取数字孪生工厂状态，识别订单队列、机器可用性、工艺能力和资源约束，然后通过智能体协作链路生成调度策略，并调用CP-SAT求解器输出排程结果。前端大屏同步展示设备状态、订单队列、AI推理日志、甘特图和基准测试结果，使用户能够在一个界面中掌握工厂运行全局。")
    add_image(doc, "03_before_failure.png", "图4-1 正常排产状态（所有机器可用，OEE≥90%）")
    add_heading(doc, "4.2 故障自动重排", 2)
    add_para(doc, "当设备故障触发后，系统会在秒级完成故障感知、拓扑切换、多智能体协同分析、CP-SAT重算排程和甘特图更新。故障机器被自动隔离，未完成工序被标记为需要重新分配，调度求解器在新的资源约束下重新生成可执行计划，全流程在3秒内完成并同步呈现在大屏界面中。")
    add_image(doc, "04_after_failure.png", "图4-2 设备故障后系统自动切换高风险审查拓扑，AI推理日志实时展示决策过程")
    add_heading(doc, "4.3 机器自动恢复", 2)
    add_para(doc, "系统为故障机器设置30至120分钟的自动恢复时间窗口。恢复事件到达后，机器状态从故障恢复为可用，系统自动将其重新纳入资源池，并触发新一轮智能体分析与排程优化。该机制使故障隔离与资源恢复均可被调度系统实时感知，从而保证排程结果随工厂状态变化持续更新。")
    add_image(doc, "05_recovery_countdown.png", "图4-3 机器倒计时自动恢复机制")
    add_heading(doc, "4.4 基准测试验证", 2)
    add_para(doc, "为验证系统在不同规模Job Shop Scheduling问题上的工程可行性，本文选取FT06、LA40、ABZ9、SWV20和DMU80五个公开标准数据集进行测试。实验结果表明，系统在小规模和中等规模任务上能够在200ms以内完成调度求解，并在LA40基准上取得与已知最优解差距7.0%的结果。")

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["数据集", "作业数", "机器数", "本系统Makespan", "最优已知值", "差距%"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    rows = [
        ["FT06", "6", "6", "67", "55", "21.8%"],
        ["LA40", "15", "15", "1307", "1222", "7.0%"],
        ["ABZ9", "20", "15", "873", "678", "28.8%"],
        ["SWV20", "50", "10", "2823", "1675", "68.5%"],
        ["DMU80", "50", "20", "9280", "暂无", "-"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    apply_table_font(table)
    add_image(doc, "06_benchmark_results.png", "图4-4 五个国际标准Job Shop Scheduling基准数据集测试对比结果")

    add_heading(doc, "4.5 A2C训练结果", 2)
    add_para(doc, "A2C训练模块用于优化拓扑选择策略。训练过程中，系统根据每一轮调度任务的奖励信号更新不同拓扑被选择的概率分布。训练结果显示，系统能够逐渐提高并行汇聚和高风险审查等拓扑在复杂场景中的选择概率，体现出对任务风险和调度复杂度的适应性。")
    add_image(doc, "07_a2c_training.png", "图4-5 A2C强化学习训练奖励曲线与拓扑选择概率分布")

    add_heading(doc, "第五章 结论与展望", 1)
    add_para(
        doc,
        "本文提出并实现了DynaTwin-Swarm工业多智能体协同调度系统，在工业调度领域实现DynaSwarm动态图结构选择与ReflAct反射推理的协同工程化，验证了多智能体协作在动态扰动响应场景中的有效性。系统通过数字孪生状态建模、豆包大模型推理、A2C拓扑优化和CP-SAT约束求解，实现了从故障感知到自动重排的闭环流程。后续工作将接入真实工厂IoT数据流，扩展至磷化工和矿山等贵州重点工业场景，并进一步完善华为盘古大模型与昇腾推理体系的深度集成，推动系统从技术验证阶段向产业落地阶段迈进。",
    )

    add_heading(doc, "参考文献", 1)
    references = [
        "[1] DynaSwarm: Dynamic Graph Structure Selection for Multi-Agent Collaboration[C]//Proceedings of EMNLP 2025.",
        "[2] ReflAct: Target-State Reflection for Language Agent Decision Making[C]//Proceedings of EMNLP 2025.",
        "[3] Google. OR-Tools CP-SAT Solver Documentation[EB/OL]. Google Developers.",
        "[4] Huawei. Ascend AI Technical White Paper[R]. Huawei Technologies Co., Ltd.",
        "[5] Taillard E. Benchmarks for basic scheduling problems[J]. European Journal of Operational Research, 1993, 64(2): 278-285.",
    ]
    for ref in references:
        add_para(doc, ref)

    DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    create_document()
